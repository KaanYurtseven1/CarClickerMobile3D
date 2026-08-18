"""
Generates SES_TASARIMI_DOKUMANI.docx from the markdown source.
Run: python generate_docx.py
"""

import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

MD_PATH = "SES_TASARIMI_DOKUMANI.md"
DOCX_PATH = "SES_TASARIMI_DOKUMANI.docx"


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_header_row(row, bg_color="1F4E79"):
    """Style header row with background color and white bold text."""
    for cell in row.cells:
        set_cell_shading(cell, bg_color)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(8)


def add_paragraph_text(doc, text, style=None, bold=False, size=None, color=None, alignment=None):
    """Add a styled paragraph."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.font.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment
    return p


def parse_md_table(lines):
    """Parse markdown table lines into headers + rows."""
    table_lines = [l for l in lines if l.strip().startswith("|")]
    if len(table_lines) < 3:
        return None, None
    
    def split_row(line):
        cells = line.strip().strip("|").split("|")
        return [c.strip() for c in cells]
    
    headers = split_row(table_lines[0])
    # Skip separator line (index 1)
    rows = [split_row(l) for l in table_lines[2:]]
    return headers, rows


def clean_md(text):
    """Remove markdown formatting artifacts."""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # bold
    text = re.sub(r'`(.*?)`', r'\1', text)  # code
    text = re.sub(r'_(.*?)_', r'\1', text)  # italic
    text = text.replace('×', '×')
    return text.strip()


def add_table_from_md(doc, md_lines, font_size=7):
    """Create a Word table from markdown table lines."""
    headers, rows = parse_md_table(md_lines)
    if not headers or not rows:
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(clean_md(h))
        run.font.size = Pt(font_size)
        run.font.bold = True

    style_header_row(table.rows[0])

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= len(headers):
                break
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            cleaned = clean_md(cell_text)
            run = p.add_run(cleaned)
            run.font.size = Pt(font_size)

            # Color status cells
            if "BAĞLI DEĞİL" in cleaned or "BAĞLI DEĞİL" in cell_text:
                set_cell_shading(cell, "FFF2CC")
                run.font.color.rgb = RGBColor(0xBF, 0x6B, 0x00)
                run.font.bold = True
            elif "BAĞLI" in cleaned and "DEĞİL" not in cleaned:
                set_cell_shading(cell, "E2EFDA")
                run.font.color.rgb = RGBColor(0x37, 0x82, 0x27)
                run.font.bold = True
            elif "MEVCUT" in cleaned:
                set_cell_shading(cell, "DAEEF3")
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                run.font.bold = True

            # Priority colors
            if "YÜKSEK" in cleaned:
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                run.font.bold = True
            elif "ORTA" in cleaned and "YÜKSEK" not in cleaned:
                run.font.color.rgb = RGBColor(0xBF, 0x8F, 0x00)
                run.font.bold = True
            elif "DÜŞÜK" in cleaned:
                run.font.color.rgb = RGBColor(0x37, 0x82, 0x27)

            # Alternate row shading (light)
            if r_idx % 2 == 1 and "BAĞLI" not in cleaned and "MEVCUT" not in cleaned:
                set_cell_shading(cell, "F5F5F5")

    return table


def build_document():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Cm(29.7)  # A4 Landscape
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # --- Default font ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(9)
    style.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # ===== COVER PAGE =====
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CAR CLICKER MOBILE 3D")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("TAM SES TASARIMI DOKÜMANI")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x4A, 0x86, 0xC8)

    doc.add_paragraph()

    meta_items = [
        ("Versiyon:", "1.0"),
        ("Tarih:", "Haziran 2025"),
        ("Hedef Kitle:", "Ses Tasarımcısı / Sound Designer"),
        ("Proje:", "CarClickerMobile3D — Mobil Idle/Clicker Araba Oyunu"),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label + " ")
        r1.font.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(value)
        r2.font.size = Pt(11)

    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    doc.add_heading("İÇİNDEKİLER", level=1)
    toc_items = [
        "1. Genel Bakış",
        "2. Sandık Sistemi (Chest System)",
        "3. Nitro Sistemi (Nitro System)",
        "4. Boost Sistemi (Boost Mode)",
        "5. Polis / Kovalamaca Sistemi (Police & Chase)",
        "6. Sistemler — Kart Efektleri",
        "7. UI / UX Sesleri",
        "8. Garaj Sahnesi (Garage Scene)",
        "9. Sinematik Sahne (Cinematic Showcase)",
        "10. Dünya Toplanabilir Öğeleri",
        "11. Arka Plan Müzikleri (Background Music)",
        "12. Bağlanmamış Sesler — Entegrasyon Bekleyenler",
        "13. Eksik Ses Fırsatları — Yeni Öneriler",
        "14. Teknik Davranış Rehberi",
        "15. Toplam Sayılar",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # Read the markdown
    with open(MD_PATH, "r", encoding="utf-8") as f:
        md_content = f.read()
    
    lines = md_content.split("\n")

    # ===== SECTION 1: GENEL BAKIŞ =====
    doc.add_heading("1. GENEL BAKIŞ", level=1)

    doc.add_heading("Proje Karakteri", level=2)
    add_paragraph_text(doc,
        "Mobil idle/clicker araba oyunu. Oyuncu ana sahne'de arabaya tıklayarak para kazanır, "
        "binalar satın alır, kart koleksiyonu oluşturur, nitro coin toplar, boost modu aktive eder, "
        "polis kovalamacalarından kaçar ve garajda arabalarını özelleştirir.",
        size=9)

    doc.add_heading("Ses Felsefesi", level=2)
    philosophies = [
        ("Ana Sahne:", "Enerjik, tatmin edici tıklama geri bildirimleri. Hızlı oynanış hissi."),
        ("Sandık Sahnesi:", "Büyülü, meraklı, ödül keşfi heyecanı."),
        ("Garaj Sahnesi:", "Showroom ambiyansı, lüks ve teknolojik."),
        ("Sinematik Sahne:", "Dramatik, prestijli, epik araba sunumu."),
        ("Polis Kovalamacası:", "Adrenalin, gerilim, kalp atışı hissi."),
        ("Genel UI:", "Hafif, temiz, rahatsız etmeyen — mobil dostu."),
    ]
    for label, desc in philosophies:
        p = doc.add_paragraph()
        r = p.add_run(f"• {label} ")
        r.font.bold = True
        r.font.size = Pt(9)
        r2 = p.add_run(desc)
        r2.font.size = Pt(9)

    doc.add_heading("Teknik Kısıtlamalar", level=2)
    constraints = [
        "Format: Mono veya Stereo, 44.1kHz, 16-bit WAV (Unity import sırasında sıkıştırılacak)",
        "Süre: One-shot efektler 0.1s–3.0s arası; loop'lar seamless olmalı",
        "Loudness: Tüm sesler -6dB peak ile normalize edilmeli (Unity tarafında volume kontrol var)",
        "Loop'lar: Seamless loop noktaları ile teslim edilmeli (crossfade point belirtilmeli)",
    ]
    for c in constraints:
        p = doc.add_paragraph()
        r = p.add_run(f"• {c}")
        r.font.size = Pt(9)

    doc.add_page_break()

    # ===== Helper: extract table block from markdown =====
    def find_table_block(start_marker, end_markers=None):
        """Find lines belonging to a markdown table section."""
        collecting = False
        result = []
        for i, line in enumerate(lines):
            if start_marker in line:
                collecting = True
                continue
            if collecting:
                if end_markers:
                    for em in end_markers:
                        if em in line:
                            return result
                if line.strip().startswith("|"):
                    result.append(line)
                elif result and not line.strip().startswith("|") and line.strip() == "":
                    # Empty line after table → table ended
                    if result:
                        return result
                elif result and line.strip().startswith("#"):
                    return result
        return result

    # ===== SECTION 2: SANDIK SİSTEMİ =====
    doc.add_heading("2. SANDIK SİSTEMİ (Chest System)", level=1)
    add_paragraph_text(doc,
        "Sandık açma sahnesi 7–8 dokunuşluk bir akıştan oluşur. Her aşamada farklı bir ses tetiklenir.",
        size=9)
    
    table_lines = find_table_block("## 2. SANDIK", ["## 3."])
    if table_lines:
        add_table_from_md(doc, table_lines)

    doc.add_page_break()

    # ===== SECTION 3: NİTRO SİSTEMİ =====
    doc.add_heading("3. NİTRO SİSTEMİ (Nitro System)", level=1)
    add_paragraph_text(doc,
        "Nitro coin'ler yolda spawn olur, oyuncu dokunarak veya mıknatıs ile toplar. "
        "Yeterince toplandığında \"Nitro Yağmuru\" başlar.",
        size=9)
    
    table_lines = find_table_block("## 3. NİTRO", ["## 4."])
    if table_lines:
        add_table_from_md(doc, table_lines)

    doc.add_page_break()

    # ===== SECTION 4: BOOST SİSTEMİ =====
    doc.add_heading("4. BOOST SİSTEMİ (Boost Mode)", level=1)
    add_paragraph_text(doc,
        "Nitro coinler boost çubuğunu doldurur → dolduğunda turbo modu aktive olur → "
        "süre boyunca güçlendirilmiş hız + ses efektleri.",
        size=9)
    
    table_lines = find_table_block("## 4. BOOST", ["## 5."])
    if table_lines:
        add_table_from_md(doc, table_lines)

    doc.add_page_break()

    # ===== SECTION 5: POLİS / KOVALAMACA =====
    doc.add_heading("5. POLİS / KOVALAMACA SİSTEMİ (Police & Chase)", level=1)
    add_paragraph_text(doc,
        "İki alt sistem: Radar (yolda beliren radarlar) ve Polis Kovalamacası (tap minigame).",
        size=9)

    # 5A
    doc.add_heading("5A. Radar Sesleri (SFXManager)", level=2)
    table_lines = find_table_block("### 5A.", ["### 5B."])
    if table_lines:
        add_table_from_md(doc, table_lines)

    # 5B
    doc.add_heading("5B. Kovalamaca Anı Sesleri (SFXManager)", level=2)
    table_lines = find_table_block("### 5B.", ["### 5C."])
    if table_lines:
        add_table_from_md(doc, table_lines)

    # 5C
    doc.add_heading("5C. Kovalamaca Katmanlı Ses Sistemi (PoliceChaseFeedbackController)", level=2)
    add_paragraph_text(doc,
        "Bu sesler ayrı AudioSource'larda çalar ve danger seviyesine göre dinamik pitch/volume değiştirir.",
        size=9)
    table_lines = find_table_block("### 5C.", ["## 6."])
    if table_lines:
        add_table_from_md(doc, table_lines)

    doc.add_page_break()

    # ===== SECTION 6: KART EFEKTLERİ =====
    doc.add_heading("6. SİSTEMLER — KART EFEKTLERİ", level=1)
    add_paragraph_text(doc,
        "Oyundaki kart efektleri çeşitli güç-up'lar sağlar. Her aktivasyon/deaktivasyon'un sesi olmalı.",
        size=9)
    
    table_lines = find_table_block("## 6. SİSTEMLER", ["## 7."])
    if table_lines:
        add_table_from_md(doc, table_lines)

    doc.add_page_break()

    # ===== SECTION 7: UI / UX =====
    doc.add_heading("7. UI / UX SESLERİ", level=1)
    add_paragraph_text(doc,
        "Tüm menü, panel, popup ve buton etkileşim sesleri.",
        size=9)

    # 7A
    doc.add_heading("7A. Temel Navigasyon", level=2)
    table_lines = find_table_block("### 7A.", ["### 7B."])
    if table_lines:
        add_table_from_md(doc, table_lines)

    # 7B
    doc.add_heading("7B. Sandık Yönetimi (Ana Ekran)", level=2)
    table_lines = find_table_block("### 7B.", ["### 7C."])
    if table_lines:
        add_table_from_md(doc, table_lines)

    # 7C
    doc.add_heading("7C. Günlük Teklifler / Ödüller", level=2)
    table_lines = find_table_block("### 7C.", ["### 7D."])
    if table_lines:
        add_table_from_md(doc, table_lines)

    # 7D
    doc.add_heading("7D. Kart & Ödül Sistemi", level=2)
    table_lines = find_table_block("### 7D.", ["### 7E."])
    if table_lines:
        add_table_from_md(doc, table_lines)

    # 7E
    doc.add_heading("7E. Genel Oyun İçi", level=2)
    table_lines = find_table_block("### 7E.", ["## 8."])
    if table_lines:
        add_table_from_md(doc, table_lines)

    doc.add_page_break()

    # ===== SECTION 8: GARAJ =====
    doc.add_heading("8. GARAJ SAHNESİ (Garage Scene)", level=1)
    add_paragraph_text(doc,
        "Garajda arabaları görüntüleme, renk/sticker/parça değiştirme, satın alma.",
        size=9)
    
    table_lines = find_table_block("## 8. GARAJ", ["## 9."])
    if table_lines:
        add_table_from_md(doc, table_lines)

    doc.add_page_break()

    # ===== SECTION 9: SİNEMATİK =====
    doc.add_heading("9. SİNEMATİK SAHNE (Cinematic Showcase)", level=1)
    add_paragraph_text(doc,
        "\"TakeTheCarScene\" — Yeni araba açıldığında dramatik sinematik sunum.",
        size=9)
    
    table_lines = find_table_block("## 9. SİNEMATİK", ["## 10."])
    if table_lines:
        add_table_from_md(doc, table_lines)

    doc.add_page_break()

    # ===== SECTION 10: DÜNYA TOPLANABILIR =====
    doc.add_heading("10. DÜNYA TOPLANABILIR ÖĞELERİ", level=1)
    
    table_lines = find_table_block("## 10. DÜNYA", ["## 11."])
    if table_lines:
        add_table_from_md(doc, table_lines)

    doc.add_page_break()

    # ===== SECTION 11: MÜZİKLER =====
    doc.add_heading("11. ARKA PLAN MÜZİKLERİ (Background Music)", level=1)
    add_paragraph_text(doc,
        "4 sahneye özel müzik parçası + crossfade geçiş sistemi.",
        size=9)
    
    table_lines = find_table_block("## 11. ARKA PLAN", ["### Müzik Sistemi"])
    if table_lines:
        add_table_from_md(doc, table_lines)

    doc.add_paragraph()
    doc.add_heading("Müzik Sistemi Özellikleri", level=2)
    music_features = [
        "Çift AudioSource Crossfade: A/B kaynakları arasında yumuşak 1.0s crossfade ile kesintisiz geçiş",
        "Duck/Restore: Boost modunda müzik %50'ye indirilir (0.3s fade), bitişte geri çıkar (0.8s fade)",
        "Sahne Otomatik Geçiş: SceneManager.sceneLoaded event'i ile sahne değiştiğinde otomatik parça değişimi",
        "Kullanıcı Volume: PlayerPrefs \"MusicVolume\" key'i ile kalıcı volume ayarı",
    ]
    for feat in music_features:
        p = doc.add_paragraph()
        r = p.add_run(f"• {feat}")
        r.font.size = Pt(9)

    doc.add_page_break()

    # ===== SECTION 12: BAĞLANMAMIŞ SESLER =====
    doc.add_heading("12. BAĞLANMAMIŞ SESLER — ENTEGRASYON BEKLEYENLER", level=1)
    add_paragraph_text(doc,
        "Aşağıdaki sesler SFXManager'da tanımlanmış ve Play metotları yazılmış ancak henüz oyun koduna "
        "entegre edilmemiş. Ses dosyaları oluşturulmalı ama programcı tarafından da bağlanmaları gerekecek.",
        size=9)
    
    table_lines = find_table_block("## 12. BAĞLANMAMIŞ", ["## 13."])
    if table_lines:
        add_table_from_md(doc, table_lines)

    doc.add_page_break()

    # ===== SECTION 13: YENİ ÖNERİLER =====
    doc.add_heading("13. EKSİK SES FIRSATLARI — YENİ ÖNERİLER", level=1)
    add_paragraph_text(doc,
        "Mevcut kod tabanında ses hook'u hiç bulunmayan ama kullanıcı deneyimini "
        "zenginleştirecek potansiyel sesler:",
        size=9)
    
    table_lines = find_table_block("## 13. EKSİK", ["## 14."])
    if table_lines:
        add_table_from_md(doc, table_lines)

    doc.add_page_break()

    # ===== SECTION 14: TEKNİK DAVRANIŞ REHBERİ =====
    doc.add_heading("14. TEKNİK DAVRANIŞ REHBERİ", level=1)

    # Volume table
    doc.add_heading("Volume Kontrolleri", level=2)
    table_lines = find_table_block("### Volume Kontrolleri", ["### Rate-Limiting"])
    if table_lines:
        add_table_from_md(doc, table_lines, font_size=8)

    # Rate-limiting table
    doc.add_heading("Rate-Limiting Kuralları", level=2)
    table_lines = find_table_block("### Rate-Limiting", ["### Pitch Variation"])
    if table_lines:
        add_table_from_md(doc, table_lines, font_size=8)

    # Pitch variation table
    doc.add_heading("Pitch Variation Kuralları", level=2)
    table_lines = find_table_block("### Pitch Variation", ["### Crossfade"])
    if table_lines:
        add_table_from_md(doc, table_lines, font_size=8)

    # Crossfade table
    doc.add_heading("Crossfade Süreleri", level=2)
    table_lines = find_table_block("### Crossfade Süreleri", ["## 15."])
    if table_lines:
        add_table_from_md(doc, table_lines, font_size=8)

    doc.add_page_break()

    # ===== SECTION 15: TOPLAM SAYILAR =====
    doc.add_heading("15. TOPLAM SAYILAR", level=1)

    doc.add_heading("Ses Dosyası İhtiyaç Özeti", level=2)
    table_lines = find_table_block("### Ses Dosyası İhtiyaç Özeti", ["### Dosya Tipi"])
    if table_lines:
        add_table_from_md(doc, table_lines, font_size=8)

    doc.add_heading("Dosya Tipi Dağılımı", level=2)
    file_types = [
        "One-shot efektler: ~85 adet",
        "Loop efektler: 7 adet (rain, boost hum, chase loop, heartbeat, siren, engine, road ambiance)",
        "Müzik parçaları: 4 adet",
        "Çoklu varyasyon: 11 adet araba dokunuş sesi",
    ]
    for ft in file_types:
        p = doc.add_paragraph()
        r = p.add_run(f"• {ft}")
        r.font.size = Pt(9)

    doc.add_heading("Öncelik Sıralaması", level=2)
    priorities = [
        ("1. YÜKSEK", "Oynanışı doğrudan etkileyen sesler (chase success/fail, turbo finger, tier advance, radar miss) — 5 ses", RGBColor(0xCC, 0x00, 0x00)),
        ("2. ORTA", "Deneyimi zenginleştiren sesler (UI hooks, popularity, garage locked, yeni öneriler) — 14 ses", RGBColor(0xBF, 0x8F, 0x00)),
        ("3. DÜŞÜK", "İnce dokunuş sesleri (garage detail, cashback, spawn efektleri) — 15 ses", RGBColor(0x37, 0x82, 0x27)),
    ]
    for label, desc, color in priorities:
        p = doc.add_paragraph()
        r = p.add_run(f"{label} — ")
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = color
        r2 = p.add_run(desc)
        r2.font.size = Pt(9)

    # ===== FOOTER NOTE =====
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "NOT: Bu doküman tamamen mevcut kod tabanı analiz edilerek oluşturulmuştur. "
        "\"✅ BAĞLI\" olarak işaretlenen sesler oyun kodunda aktif olarak tetiklenmektedir. "
        "\"⚠️ BAĞLI DEĞİL\" olarak işaretlenenler için Play metotları mevcuttur ancak oyun akışına "
        "entegrasyonları programcı tarafından yapılacaktır."
    )
    r.font.size = Pt(8)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ===== SAVE =====
    doc.save(DOCX_PATH)
    print(f"✅ Word belgesi oluşturuldu: {DOCX_PATH}")


if __name__ == "__main__":
    build_document()
